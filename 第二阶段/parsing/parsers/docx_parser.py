"""DOCX 文本和表格解析器。"""

from pathlib import Path

from 第二阶段.parsing.base import BaseParser, ParserDependencyError, ParserError
from 第二阶段.schemas.models import ParsedDocument


class DocxParser(BaseParser):
    file_type = "docx"

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = self._validate_path(file_path)
        try:
            from docx import Document
        except ImportError as exc:
            raise ParserDependencyError("DOCX 解析需要安装 python-docx") from exc
        try:
            document = Document(str(path))
            blocks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        blocks.append("\t".join(cells))
        except Exception as exc:
            raise ParserError(f"DOCX 解析失败：{path}") from exc
        return self._build_document(
            path,
            "\n\n".join(blocks),
            metadata={"paragraph_count": len(document.paragraphs), "table_count": len(document.tables)},
        )

